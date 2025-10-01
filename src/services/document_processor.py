import logging
import asyncio
import aiofiles
import os
import re
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import pypdf
import docx
from datetime import datetime
import uuid

from ..models.document_models import Document, DocumentChunk, DocumentMetadata, DocumentType, DocumentStatus

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Advanced document processing service for RAG pipeline with optimized chunking"""
    
    def __init__(self, upload_dir: str = "uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
        
        # OPTIMIZED Processing settings for better chat responses
        self.max_chunk_size = 500   # Reduced from 1000 - shorter chunks = shorter responses
        self.chunk_overlap = 100    # Reduced from 200 
        self.min_chunk_size = 50    # Reduced from 100 - allow smaller meaningful chunks
        
        # Store processed chunks in memory for retrieval
        self.document_chunks = {}  # document_id -> List[DocumentChunk]
        
        logger.info(f"Document Processor initialized - Upload dir: {self.upload_dir}")
        logger.info(f"Chunking settings: max={self.max_chunk_size}, overlap={self.chunk_overlap}, min={self.min_chunk_size}")
    
    async def process_uploaded_file(self, file_content: bytes, filename: str, 
                                  file_type: DocumentType, description: Optional[str] = None) -> Document:
        """
        Process uploaded file and return document object
        """
        try:
            # Generate unique filename and save file
            document_id = str(uuid.uuid4())
            safe_filename = self._sanitize_filename(filename)
            file_path = self.upload_dir / f"{document_id}_{safe_filename}"
            
            # Save file to disk
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(file_content)
            
            # Create document object
            document = Document(
                id=document_id,
                filename=safe_filename,
                original_filename=filename,
                file_path=str(file_path),
                file_size=len(file_content),
                file_type=file_type,
                status=DocumentStatus.PROCESSING,
                description=description
            )
            
            logger.info(f"File saved: {filename} -> {file_path}")
            
            # Process document content
            await self._process_document_content(document)
            
            return document
            
        except Exception as e:
            logger.error(f"Error processing uploaded file {filename}: {str(e)}")
            # Create failed document record
            return Document(
                id=str(uuid.uuid4()),
                filename=filename,
                original_filename=filename,
                file_path="",
                file_size=len(file_content),
                file_type=file_type,
                status=DocumentStatus.FAILED,
                error_message=str(e),
                description=description
            )
    
    async def _process_document_content(self, document: Document) -> None:
        """
        Extract and process document content
        """
        start_time = datetime.now()
        
        try:
            # Extract text content
            text_content, metadata = await self._extract_content(document.file_path, document.file_type)
            
            # Update document metadata
            document.metadata = metadata
            
            # Create optimized chunks
            chunks = self._create_intelligent_chunks(text_content, document.id)
            document.chunks_count = len(chunks)
            
            # IMPORTANT: Store chunks in memory for later retrieval
            self.document_chunks[document.id] = chunks
            
            # Update status and timing
            document.status = DocumentStatus.COMPLETED
            document.processing_time = (datetime.now() - start_time).total_seconds()
            document.updated_at = datetime.now()
            
            logger.info(f"Document processed successfully: {document.filename} - {len(chunks)} chunks stored")
            
        except Exception as e:
            document.status = DocumentStatus.FAILED
            document.error_message = str(e)
            document.updated_at = datetime.now()
            logger.error(f"Error processing document content: {str(e)}")
    
    async def get_document_chunks(self, document_id: str) -> List[Dict[str, Any]]:
        """
        Get processed chunks for a document - THIS IS THE KEY METHOD!
        """
        try:
            # First try to get from memory
            if document_id in self.document_chunks:
                chunks = self.document_chunks[document_id]
                logger.info(f"Retrieved {len(chunks)} chunks from memory for document {document_id}")
                
                # Convert DocumentChunk objects to dictionaries
                chunk_dicts = []
                for chunk in chunks:
                    chunk_dict = {
                        'content': chunk.content,
                        'chunk_index': chunk.chunk_index,
                        'metadata': chunk.metadata or {'document_id': document_id}
                    }
                    chunk_dicts.append(chunk_dict)
                
                return chunk_dicts
            
            # If not in memory, try to reprocess from file
            logger.info(f"Chunks not in memory, attempting to reprocess document {document_id}")
            
            # Find the document file
            for file_path in self.upload_dir.glob(f"{document_id}_*"):
                if file_path.exists():
                    # Determine file type and reprocess
                    file_type = self.validate_file_type(file_path.name)
                    if file_type:
                        try:
                            text_content, _ = await self._extract_content(str(file_path), file_type)
                            chunks = self._create_intelligent_chunks(text_content, document_id)
                            
                            # Store back in memory
                            self.document_chunks[document_id] = chunks
                            
                            # Convert to dictionaries
                            chunk_dicts = []
                            for chunk in chunks:
                                chunk_dict = {
                                    'content': chunk.content,
                                    'chunk_index': chunk.chunk_index,
                                    'metadata': chunk.metadata or {'document_id': document_id}
                                }
                                chunk_dicts.append(chunk_dict)
                            
                            logger.info(f"Reprocessed and retrieved {len(chunk_dicts)} chunks for document {document_id}")
                            return chunk_dicts
                            
                        except Exception as e:
                            logger.error(f"Error reprocessing document for chunks: {str(e)}")
            
            # If all else fails, return empty list
            logger.warning(f"No chunks found for document {document_id}")
            return []
                        
        except Exception as e:
            logger.error(f"Error getting document chunks: {str(e)}")
            return []
    
    def _create_intelligent_chunks(self, text_content: str, document_id: str) -> List[DocumentChunk]:
        """
        Create intelligent chunks with better size control for optimal chat responses
        """
        chunks = []
        
        if not text_content or not text_content.strip():
            logger.warning(f"No text content to chunk for document {document_id}")
            return chunks
        
        logger.info(f"Creating chunks for document {document_id} - text length: {len(text_content)} chars")
        
        # Clean the text first
        text_content = self._clean_text(text_content)
        
        # Strategy 1: Try to find natural sections
        sections = self._extract_sections(text_content)
        
        # Strategy 2: If no sections, split by paragraphs
        if not sections or len(sections) <= 2:
            logger.info("No clear sections found, splitting by paragraphs")
            paragraphs = [p.strip() for p in text_content.split('\n\n') if p.strip() and len(p.strip()) > 20]
            if paragraphs:
                sections = paragraphs
            else:
                # Fallback: split by sentences
                sentences = [s.strip() + '.' for s in text_content.split('.') if s.strip() and len(s.strip()) > 10]
                sections = self._group_sentences(sentences)
        
        logger.info(f"Split text into {len(sections)} initial sections")
        
        # Create chunks from sections
        chunk_index = 0
        
        for i, section in enumerate(sections):
            if not section or len(section.strip()) < self.min_chunk_size:
                continue
                
            # If section is too large, split it further
            if len(section) > self.max_chunk_size:
                sub_chunks = self._split_large_section(section, document_id, chunk_index)
                chunks.extend(sub_chunks)
                chunk_index += len(sub_chunks)
            else:
                # Section is good size, create chunk
                chunks.append(DocumentChunk(
                    content=section.strip(),
                    chunk_index=chunk_index,
                    metadata={
                        'document_id': document_id,
                        'token_count': len(section.split()),
                        'char_count': len(section),
                        'chunk_type': 'section',
                        'section_index': i
                    }
                ))
                chunk_index += 1
        
        # Quality check and merging of very small chunks
        chunks = self._optimize_chunks(chunks, document_id)
        
        avg_size = sum(len(c.content) for c in chunks) // max(len(chunks), 1)
        logger.info(f"Created {len(chunks)} optimized chunks for document {document_id} (avg size: {avg_size} chars)")
        
        return chunks
    
    def _extract_sections(self, text: str) -> List[str]:
        """
        Extract natural sections from text based on headers and structure
        """
        sections = []
        
        # Patterns for section headers (academic papers, reports, etc.)
        section_patterns = [
            r'\n\s*(?:Abstract|Introduction|Background|Literature Review|Methodology|Method|Methods|Results|Discussion|Conclusion|Conclusions|References|Bibliography)\s*\n',
            r'\n\s*\d+\.?\s+[A-Z][^.\n]{5,50}\s*\n',  # "1. Introduction" style
            r'\n\s*[A-Z][A-Z\s]{3,30}[A-Z]\s*\n',     # "INTRODUCTION" style
            r'\n\s*[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s*\n'  # "Introduction Method" style
        ]
        
        split_positions = [0]  # Start position
        
        for pattern in section_patterns:
            matches = re.finditer(pattern, text, re.MULTILINE)
            for match in matches:
                split_positions.append(match.start())
        
        # Sort and deduplicate positions
        split_positions = sorted(set(split_positions))
        split_positions.append(len(text))  # End position
        
        # Extract sections
        for i in range(len(split_positions) - 1):
            start = split_positions[i]
            end = split_positions[i + 1]
            section = text[start:end].strip()
            
            if section and len(section) > 50:  # Only meaningful sections
                sections.append(section)
        
        return sections
    
    def _group_sentences(self, sentences: List[str]) -> List[str]:
        """
        Group sentences into reasonable chunks
        """
        groups = []
        current_group = ""
        
        for sentence in sentences:
            if len(current_group) + len(sentence) > self.max_chunk_size and current_group:
                groups.append(current_group.strip())
                current_group = sentence
            else:
                current_group += " " + sentence if current_group else sentence
        
        if current_group:
            groups.append(current_group.strip())
        
        return groups
    
    def _split_large_section(self, section: str, document_id: str, start_index: int) -> List[DocumentChunk]:
        """
        Split a large section into smaller chunks with overlap
        """
        chunks = []
        
        # Try splitting by paragraphs first
        paragraphs = [p.strip() for p in section.split('\n\n') if p.strip()]
        
        if not paragraphs:
            # Split by sentences
            sentences = [s.strip() + '.' for s in section.split('.') if s.strip()]
            paragraphs = self._group_sentences(sentences)
        
        current_chunk = ""
        chunk_index = start_index
        
        for para in paragraphs:
            if len(current_chunk) + len(para) > self.max_chunk_size and current_chunk:
                # Finalize current chunk
                if len(current_chunk.strip()) >= self.min_chunk_size:
                    chunks.append(DocumentChunk(
                        content=current_chunk.strip(),
                        chunk_index=chunk_index,
                        metadata={
                            'document_id': document_id,
                            'token_count': len(current_chunk.split()),
                            'char_count': len(current_chunk),
                            'chunk_type': 'split_large'
                        }
                    ))
                    chunk_index += 1
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk, self.chunk_overlap)
                current_chunk = overlap_text + "\n\n" + para if overlap_text else para
            else:
                current_chunk += "\n\n" + para if current_chunk else para
        
        # Add final chunk
        if current_chunk.strip() and len(current_chunk.strip()) >= self.min_chunk_size:
            chunks.append(DocumentChunk(
                content=current_chunk.strip(),
                chunk_index=chunk_index,
                metadata={
                    'document_id': document_id,
                    'token_count': len(current_chunk.split()),
                    'char_count': len(current_chunk),
                    'chunk_type': 'split_large_final'
                }
            ))
        
        return chunks
    
    def _optimize_chunks(self, chunks: List[DocumentChunk], document_id: str) -> List[DocumentChunk]:
        """
        Optimize chunks by merging very small ones and ensuring quality
        """
        if not chunks:
            return chunks
        
        optimized = []
        current_chunk = chunks[0]
        
        for next_chunk in chunks[1:]:
            # If current chunk is too small, try to merge with next
            if (len(current_chunk.content) < self.min_chunk_size * 2 and 
                len(current_chunk.content) + len(next_chunk.content) <= self.max_chunk_size):
                
                # Merge chunks
                merged_content = current_chunk.content + "\n\n" + next_chunk.content
                current_chunk = DocumentChunk(
                    content=merged_content,
                    chunk_index=current_chunk.chunk_index,
                    metadata={
                        'document_id': document_id,
                        'token_count': len(merged_content.split()),
                        'char_count': len(merged_content),
                        'chunk_type': 'merged'
                    }
                )
            else:
                # Keep current chunk and move to next
                optimized.append(current_chunk)
                current_chunk = next_chunk
        
        # Don't forget the last chunk
        optimized.append(current_chunk)
        
        return optimized
    
    async def _extract_content(self, file_path: str, file_type: DocumentType) -> Tuple[str, DocumentMetadata]:
        """
        Extract text content and metadata from different file types
        """
        if file_type == DocumentType.PDF:
            return await self._extract_pdf_content(file_path)
        elif file_type == DocumentType.DOCX:
            return await self._extract_docx_content(file_path)
        elif file_type == DocumentType.TXT:
            return await self._extract_text_content(file_path)
        elif file_type == DocumentType.MD:
            return await self._extract_text_content(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    async def _extract_pdf_content(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Extract content from PDF file using modern pypdf"""
        try:
            text_content = ""
            metadata = DocumentMetadata()
            
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                # Extract metadata
                if pdf_reader.metadata:
                    metadata.title = pdf_reader.metadata.get('/Title')
                    metadata.author = pdf_reader.metadata.get('/Author')
                    metadata.subject = pdf_reader.metadata.get('/Subject')
                    metadata.creator = pdf_reader.metadata.get('/Creator')
                    
                    # Handle creation date
                    creation_date = pdf_reader.metadata.get('/CreationDate')
                    if creation_date:
                        try:
                            date_str = str(creation_date).replace('D:', '').split('+')[0].split('-')[0]
                            if len(date_str) >= 8:
                                metadata.creation_date = datetime.strptime(date_str[:8], '%Y%m%d')
                        except Exception as date_error:
                            logger.warning(f"Could not parse PDF creation date: {date_error}")
                
                # Extract text from all pages
                metadata.pages = len(pdf_reader.pages)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_content += f"\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        continue
            
            # Clean and normalize text
            text_content = self._clean_text(text_content)
            metadata.word_count = len(text_content.split()) if text_content else 0
            
            logger.info(f"Extracted {len(text_content)} characters from PDF")
            return text_content, metadata
            
        except Exception as e:
            logger.error(f"Error extracting PDF content: {str(e)}")
            raise
    
    async def _extract_docx_content(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Extract content from DOCX file"""
        try:
            doc = docx.Document(file_path)
            
            text_content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            metadata = DocumentMetadata()
            try:
                core_props = doc.core_properties
                metadata.title = core_props.title
                metadata.author = core_props.author
                metadata.subject = core_props.subject
                metadata.creator = core_props.last_modified_by
                metadata.creation_date = core_props.created
                metadata.pages = len(doc.paragraphs)
            except Exception as meta_error:
                logger.warning(f"Could not extract DOCX metadata: {meta_error}")
            
            text_content = self._clean_text(text_content)
            metadata.word_count = len(text_content.split()) if text_content else 0
            
            return text_content, metadata
            
        except Exception as e:
            logger.error(f"Error extracting DOCX content: {str(e)}")
            raise
    
    async def _extract_text_content(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Extract content from plain text file"""
        try:
            encodings = ['utf-8', 'utf-16', 'ascii', 'latin-1']
            text_content = ""
            
            for encoding in encodings:
                try:
                    async with aiofiles.open(file_path, 'r', encoding=encoding) as f:
                        text_content = await f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if not text_content:
                raise ValueError("Could not decode text file with any supported encoding")
            
            metadata = DocumentMetadata()
            metadata.word_count = len(text_content.split()) if text_content else 0
            metadata.pages = text_content.count('\n') // 50 + 1
            
            text_content = self._clean_text(text_content)
            return text_content, metadata
            
        except Exception as e:
            logger.error(f"Error extracting text content: {str(e)}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace but preserve paragraph breaks
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Multiple newlines to double newline
        
        # Remove page markers and artifacts
        text = re.sub(r'--- Page \d+ ---', '', text)
        text = re.sub(r'\f', '\n', text)  # Form feed to newline
        
        # Fix common PDF extraction issues
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # Missing spaces
        text = re.sub(r'(\.)([A-Z])', r'\1 \2', text)     # Sentence boundaries
        
        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f]', '', text)
        
        return text.strip()
    
    def _get_overlap_text(self, text: str, overlap_size: int) -> str:
        """Get overlap text from the end of current chunk"""
        if not text or len(text) <= overlap_size:
            return text
        
        overlap_text = text[-overlap_size:]
        sentence_end = overlap_text.rfind('.')
        
        if sentence_end > overlap_size // 2:
            return overlap_text[sentence_end + 1:].strip()
        else:
            return overlap_text.strip()
    
    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename for safe storage"""
        if not filename:
            return "unnamed_file"
        
        safe_name = re.sub(r'[^\w\-_\.]', '_', filename)
        safe_name = re.sub(r'_+', '_', safe_name)
        
        if len(safe_name) > 100:
            name, ext = os.path.splitext(safe_name)
            safe_name = name[:95] + ext
        
        return safe_name or "unnamed_file"
    
    async def delete_document(self, document: Document) -> bool:
        """Delete document file and cleanup"""
        try:
            # Remove from memory
            if document.id in self.document_chunks:
                del self.document_chunks[document.id]
                logger.info(f"Removed chunks from memory for document {document.id}")
            
            # Remove file
            if document.file_path and os.path.exists(document.file_path):
                os.remove(document.file_path)
                logger.info(f"Deleted document file: {document.file_path}")
                return True
            return False
        except Exception as e:
            logger.error(f"Error deleting document file: {str(e)}")
            return False
    
    def get_supported_file_types(self) -> List[str]:
        """Get list of supported file types"""
        return [doc_type.value for doc_type in DocumentType]
    
    def validate_file_type(self, filename: str) -> Optional[DocumentType]:
        """Validate and determine file type from filename"""
        if not filename:
            return None
        
        ext = Path(filename).suffix.lower()
        type_mapping = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.txt': DocumentType.TXT,
            '.md': DocumentType.MD
        }
        return type_mapping.get(ext)
    
    def is_file_type_supported(self, filename: str) -> bool:
        """Check if file type is supported"""
        return self.validate_file_type(filename) is not None
    
    async def get_document_preview(self, document: Document, max_length: int = 500) -> str:
        """Get a preview of document content"""
        try:
            if document.status != DocumentStatus.COMPLETED:
                return f"Document is {document.status.value}"
            
            if not document.file_path or not os.path.exists(document.file_path):
                return "Document file not found"
            
            text_content, _ = await self._extract_content(document.file_path, document.file_type)
            
            if not text_content:
                return "Document appears to be empty or could not be processed"
            
            if len(text_content) <= max_length:
                return text_content
            else:
                return text_content[:max_length] + "..."
                
        except Exception as e:
            logger.error(f"Error getting document preview: {str(e)}")
            return f"Error loading preview: {str(e)}"
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get processing statistics"""
        total_chunks = sum(len(chunks) for chunks in self.document_chunks.values())
        
        return {
            "max_chunk_size": self.max_chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "min_chunk_size": self.min_chunk_size,
            "upload_directory": str(self.upload_dir),
            "supported_file_types": self.get_supported_file_types(),
            "documents_in_memory": len(self.document_chunks),
            "total_chunks_stored": total_chunks,
            "avg_chunks_per_document": total_chunks / max(len(self.document_chunks), 1)
        }
    
    def clear_document_cache(self, document_id: Optional[str] = None):
        """Clear document chunks from memory"""
        if document_id:
            if document_id in self.document_chunks:
                del self.document_chunks[document_id]
                logger.info(f"Cleared chunks cache for document {document_id}")
        else:
            self.document_chunks.clear()
            logger.info("Cleared all document chunks from cache")
